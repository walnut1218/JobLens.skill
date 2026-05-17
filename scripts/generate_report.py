#!/usr/bin/env python3
"""BOSS直聘职位分析报告生成脚本 - 深度分析版

报告结构对应 references/analysis_prompt.md 的深度分析模板：
  第一部分：全景总结（市场画像 / 能力结构图谱 / 行动路径建议）
  第二部分：逐条深度分析（深度解读 / 能力谱系 / 实现路径 / 作品集方案 / 面试模拟 / 简历匹配）

用法：
  uv run python scripts/generate_report.py --city 上海 --keyword "用户研究实习生"
  uv run python scripts/generate_report.py --city 上海 --keyword "用户研究实习生" --resume user_resume.txt
"""
import json, re, sys
from datetime import datetime
from pathlib import Path

_script_dir = Path(__file__).resolve().parent
if str(_script_dir) not in sys.path:
    sys.path.insert(0, str(_script_dir))

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from _resume_matcher import score_and_sort, load_resume


# ═══════════════════════════════════════════════════════════════
# 样式工具
# ═══════════════════════════════════════════════════════════════

def _set_font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

def _para(doc, text, bold=False, size=11, color=None, space=Pt(6), indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = indent
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = space
    return p

def _bullet(doc, text, size=11, color=None):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    _set_font(r, size=size, color=color)

def _quote_block(doc, text):
    """缩进引用块（用于 JD 原文）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_font(r, size=10, color=(60, 60, 60))

def _section(doc, title):
    _para(doc, "", space=Pt(8))
    _heading(doc, title, level=2)

def _subsection(doc, title):
    _heading(doc, title, level=3)

def _divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("─" * 50)
    _set_font(r, size=8, color=(180, 180, 180))


# ═══════════════════════════════════════════════════════════════
# 第一部分：全景总结
# ═══════════════════════════════════════════════════════════════

def add_overview(doc, jobs):
    """整体总结——市场画像 / 能力结构图谱 / 行动路径"""
    _heading(doc, "第一部分：全景总结", level=0)

    total = len(jobs)

    # 统计
    industries = {}
    cities = {}
    salary_ranges = {}
    exp_levels = {}
    companies = set()
    for j in jobs:
        ind = j.get("industry", "").strip()
        if ind:
            industries[ind] = industries.get(ind, 0) + 1
        c = j.get("city", "").strip()
        if c:
            cities[c] = cities.get(c, 0) + 1
        s = j.get("salary", "").strip()
        if s:
            salary_ranges[s] = salary_ranges.get(s, 0) + 1
        e = j.get("exp", "").strip()
        if e:
            exp_levels[e] = exp_levels.get(e, 0) + 1
        comp = j.get("company", "").strip()
        if comp:
            companies.add(comp)

    has_desc_count = sum(1 for j in jobs if j.get("description", "").strip())

    _section(doc, "📊 市场画像")

    _para(doc, f"共采集 {total} 个职位，来自 {len(companies)} 家公司，{has_desc_count}/{total} 有完整描述。")

    _subsection(doc, "行业分布")
    for ind, cnt in sorted(industries.items(), key=lambda x: -x[1])[:8]:
        _bullet(doc, f"{ind}：{cnt} 个")

    _subsection(doc, "薪资区间")
    for s, cnt in sorted(salary_ranges.items(), key=lambda x: -x[1]):
        _bullet(doc, f"{s}：{cnt} 个")

    _subsection(doc, "经验要求")
    for e, cnt in sorted(exp_levels.items(), key=lambda x: -x[1]):
        _bullet(doc, f"{e}：{cnt} 个")

    _subsection(doc, "公司规模与融资阶段分布")
    scales = {}
    stages = {}
    for j in jobs:
        sc = j.get("scale", "").strip()
        if sc:
            scales[sc] = scales.get(sc, 0) + 1
        st = j.get("stage", "").strip()
        if st:
            stages[st] = stages.get(st, 0) + 1
    if scales:
        _para(doc, "规模：", bold=True, size=10)
        for k, v in sorted(scales.items(), key=lambda x: -x[1]):
            _bullet(doc, f"{k}：{v} 个")
    if stages:
        _para(doc, "融资：", bold=True, size=10)
        for k, v in sorted(stages.items(), key=lambda x: -x[1]):
            _bullet(doc, f"{k}：{v} 个")

    _section(doc, "🎯 能力结构图谱")

    # 高频技能提取
    skill_freq = {}
    skill_patterns = [
        "Python", "SPSS", "SQL", "R", "Excel", "Tableau", "PowerBI",
        "问卷", "访谈", "可用性测试", "焦点小组", "田野调查",
        "数据分析", "定量分析", "定性分析", "统计分析",
        "用户画像", "用户旅程", "竞品分析", "A/B测试", "AB测试",
        "NLP", "自然语言处理", "机器学习", "深度学习",
        "Figma", "Axure", "Sketch", "用户研究", "用户调研",
        "项目管理", "需求分析", "产品设计", "交互设计",
        "沟通", "协作", "报告", "汇报", "英语", "日语",
    ]
    for j in jobs:
        desc = j.get("description", "") + j.get("name", "")
        for sp in skill_patterns:
            if sp.lower() in desc.lower():
                skill_freq[sp] = skill_freq.get(sp, 0) + 1

    if skill_freq:
        _para(doc, "硬通货（高频出现，没商量余地）：", bold=True)
        for sk, cnt in sorted(skill_freq.items(), key=lambda x: -x[1])[:8]:
            _bullet(doc, f"{sk}（{cnt}/{total} 个岗位提及）")

    _subsection(doc, "方向的真实信号")
    # 实习生比例
    intern_count = sum(1 for j in jobs if "实习" in j.get("name", ""))
    _bullet(doc, f"实习生岗位占比：{intern_count}/{total}")
    _bullet(doc, f"注：以下为逐条深度分析。每一条都不是翻译JD，而是尝试解读JD背后没说出来的信息。")


# ═══════════════════════════════════════════════════════════════
# 第二部分：逐条深度分析
# ═══════════════════════════════════════════════════════════════

def add_job_detail(doc, job, idx, has_resume=False):
    """逐条深度分析——对应 analysis_prompt.md 第二部分模板"""

    name = job.get("name", "未知")
    company = job.get("company", "未知")
    salary = job.get("salary", "面议")
    city = job.get("city", "")
    exp = job.get("exp", "")
    edu = job.get("edu", "")
    industry = job.get("industry", "")
    scale = job.get("scale", "")
    stage = job.get("stage", "")
    description = job.get("description", "")
    url = job.get("url", "")
    has_desc = bool(description and len(description.strip()) > 20)
    match_score = job.get("_match_score", -1)
    is_top = match_score >= 60

    # 岗位标题
    score_badge = ""
    if match_score >= 0:
        score_badge = f"  🏆 匹配度 {match_score}/100" if is_top else f"  匹配度 {match_score}/100"
    _heading(doc, f"职位 {idx}：{name} — {company}{score_badge}", level=1)

    # 快速标签
    tags = []
    if city: tags.append(f"📍 {city}")
    if salary: tags.append(f"💰 {salary}")
    if exp: tags.append(f"⏱ {exp}")
    if edu: tags.append(f"🎓 {edu}")
    _para(doc, " | ".join(tags), size=10, color=(90, 90, 90))

    _bullet(doc, f"🏢 {company} | {industry} | {scale} | {stage}")
    _bullet(doc, f"🔗 {url}")

    # ── 职位描述原文 ──
    _subsection(doc, "📋 职位描述原文")
    if has_desc:
        _quote_block(doc, description)
    else:
        _para(doc, "⚠️ 该职位无详细描述，以下分析基于岗位名称和标签推测，建议获取完整 JD 后复核。",
              size=10, color=(180, 80, 80))

    # ── 深度解读 ──
    _subsection(doc, "🧠 这份 JD 到底在说什么")
    _para(doc, "【⚠️ 本条需 AI 填充：从 JD 措辞推断——团队状态、岗位定位、JD 没写但真实存在的问题】",
          size=10, color=(120, 120, 120))

    # ── 真正需要的能力 ──
    _subsection(doc, "🎯 真正需要的能力（门槛级 / 竞争级 / 溢价级）")
    _para(doc, "🔸 门槛级（没有这些，简历直接挂）：", bold=True)
    _para(doc, "【⚠️ AI 填充：2-3 条，不是技能名词是场景判断】", size=10, color=(120, 120, 120))
    _para(doc, "🔸 竞争级（具备这些，大概率拿下面试）：", bold=True)
    _para(doc, "【⚠️ AI 填充：2-3 条】", size=10, color=(120, 120, 120))
    _para(doc, "🔸 溢价级（有这些，薪资可以往上谈）：", bold=True)
    _para(doc, "【⚠️ AI 填充：1-2 条】", size=10, color=(120, 120, 120))

    # ── 技能实现路径 ──
    _subsection(doc, "📋 技能实现路径")
    _para(doc, "【⚠️ AI 填充：每项核心技能 → 够用程度 → 最快上手路径（具体到天）→ 加分操作 → 避坑】",
          size=10, color=(120, 120, 120))

    # ── 作品集落地方案 ──
    _subsection(doc, "🛠️ 作品集落地方案")
    _para(doc, "【⚠️ AI 填充：项目目标 → 具体内容 → 数据来源 → 技术栈 → 预期产出 → 完成时间 → 对应JD哪几条 → 为什么选这个不选别的】",
          size=10, color=(120, 120, 120))

    # ── 面试模拟 ──
    _subsection(doc, "💬 面试模拟")
    _para(doc, "【⚠️ AI 填充：2-3 个这家公司特有的场景题 → 面试官真实意图 → 回答核心逻辑 → 加分操作】",
          size=10, color=(120, 120, 120))

    # ── 简历匹配度 ──
    if has_resume and match_score >= 0:
        _subsection(doc, f"📊 简历匹配度评估（{match_score}/100）")
        _para(doc, "✅ 直接命中的点：", bold=True)
        _para(doc, "【⚠️ AI 填充：3 条】", size=10, color=(120, 120, 120))
        _para(doc, "❌ 硬伤：", bold=True)
        _para(doc, "【⚠️ AI 填充：1-2 条需要马上解决的差距】", size=10, color=(120, 120, 120))
        _para(doc, "📝 一周突击方案：", bold=True)
        _para(doc, "【⚠️ AI 填充：具体行动，不是空话】", size=10, color=(120, 120, 120))

    _divider(doc)


# ═══════════════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════════════

def generate_report(jobs, output_path, resume_path="user_resume.txt"):
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ── 预处理 ──
    resume = load_resume(resume_path)
    has_resume = resume is not None

    # 简历评分排序
    already_scored = any(j.get("_match_score", -1) >= 0 for j in jobs)
    if not already_scored:
        jobs = score_and_sort(jobs, resume_path)

    # ── 封面 ──
    first = jobs[0] if jobs else {}
    kw = (first.get("name", "未知").split("（")[0])[:20] if first else "未知"
    ct = first.get("city", "未知")

    title = f"{ct} · {kw} · 职位深度分析报告"
    if has_resume:
        title += "（按匹配度排序）"
    _heading(doc, title, level=0)

    info = f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')} | 共 {len(jobs)} 个职位"
    if has_resume:
        top_n = sum(1 for j in jobs if j.get("_match_score", 0) >= 60)
        info += f" | 高匹配（≥60分）：{top_n} 个"
    _para(doc, info, size=10, color=(110, 110, 110), space=Pt(12))

    # Top 匹配预览
    if has_resume:
        top5 = [j for j in jobs if j.get("_match_score", 0) >= 60][:5]
        if top5:
            _heading(doc, "🏆 最佳匹配 Top 5", level=2)
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = 'Light Grid Accent 1'
            for cell, txt in zip(tbl.rows[0].cells, ["排名", "公司", "职位", "匹配度"]):
                r = cell.paragraphs[0].add_run(txt)
                _set_font(r, bold=True, size=10)
            for i, j in enumerate(top5, 1):
                row = tbl.add_row().cells
                for cell, txt in zip(row, [str(i), j.get("company", "")[:18],
                                          j.get("name", "")[:24],
                                          f"{j.get('_match_score', 0)}/100"]):
                    r = cell.paragraphs[0].add_run(txt)
                    _set_font(r, size=10)

    # ── 第一部分：全景总结 ──
    if len(jobs) > 10:
        add_overview(doc, jobs)
        _divider(doc)
        _heading(doc, "第二部分：逐条深度分析", level=0)

    # ── 第二部分：逐条分析 ──
    for i, job in enumerate(jobs, 1):
        add_job_detail(doc, job, i, has_resume)
        if i % 3 == 0 and i < len(jobs):
            doc.add_page_break()

    doc.save(output_path)
    print(f"✅ 报告已生成：{output_path}", file=sys.stderr)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="生成深度分析报告")
    parser.add_argument("--keyword", default="", help="搜索关键词")
    parser.add_argument("--city", default="", help="城市")
    parser.add_argument("--resume", default="user_resume.txt", help="简历文件路径")
    args = parser.parse_args()

    with open("jobs_data.json", "r", encoding="utf-8") as f:
        jobs = json.load(f)

    print(f"📊 加载 {len(jobs)} 个职位", file=sys.stderr)

    ct = args.city or (jobs[0].get("city", "未知") if jobs else "未知")
    kw = args.keyword or "未知"
    out = f"分析报告_{ct}_{kw}_{datetime.now().strftime('%Y%m%d')}.docx"
    out = re.sub(r'[\\/:*?"<>|]', '_', out)

    generate_report(jobs, out, args.resume)
    print(f"📄 {out}", file=sys.stderr)


if __name__ == "__main__":
    main()